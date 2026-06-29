#!/usr/bin/env python3
"""Build a compact one-page-style Chinese PM resume DOCX from structured JSON."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT_CN = "Microsoft YaHei"
FONT_FALLBACK = "PingFang SC"
ACCENT = RGBColor(31, 78, 121)
TEXT = RGBColor(30, 30, 30)

LAYOUTS = {
    "ultra": {"margin": 0.40, "normal": 8.5, "name": 15.5, "info": 8.4, "heading": 10.0, "section_before": 1, "line": 0.95, "bullet": 8.3, "small": 8.1, "entry": 8.5, "gap": 0},
    "tight": {"margin": 0.45, "normal": 8.8, "name": 16.0, "info": 8.7, "heading": 10.2, "section_before": 2, "line": 1.0, "bullet": 8.6, "small": 8.3, "entry": 8.8, "gap": 0.3},
    "normal": {"margin": 0.52, "normal": 9.2, "name": 17.0, "info": 9.0, "heading": 10.5, "section_before": 3, "line": 1.02, "bullet": 8.9, "small": 8.6, "entry": 9.2, "gap": 0.6},
    "roomy": {"margin": 0.62, "normal": 9.8, "name": 18.0, "info": 9.5, "heading": 11.0, "section_before": 5, "line": 1.08, "bullet": 9.5, "small": 9.0, "entry": 9.8, "gap": 1.2},
    "expanded": {"margin": 0.72, "normal": 10.3, "name": 19.0, "info": 9.9, "heading": 11.5, "section_before": 7, "line": 1.15, "bullet": 10.0, "small": 9.4, "entry": 10.3, "gap": 1.8},
    "fill1": {"margin": 0.78, "normal": 10.6, "name": 19.5, "info": 10.1, "heading": 11.8, "section_before": 8, "line": 1.18, "bullet": 10.3, "small": 9.7, "entry": 10.5, "gap": 2.2},
    "fill2": {"margin": 0.82, "normal": 10.9, "name": 20.0, "info": 10.4, "heading": 12.0, "section_before": 9, "line": 1.20, "bullet": 10.5, "small": 9.9, "entry": 10.7, "gap": 2.6},
    "fill3": {"margin": 0.86, "normal": 11.1, "name": 20.3, "info": 10.6, "heading": 12.2, "section_before": 10, "line": 1.22, "bullet": 10.7, "small": 10.1, "entry": 10.9, "gap": 3.0},
    "fill4": {"margin": 0.90, "normal": 11.3, "name": 20.6, "info": 10.8, "heading": 12.4, "section_before": 11, "line": 1.24, "bullet": 10.9, "small": 10.3, "entry": 11.1, "gap": 3.4},
    "filled": {"margin": 0.94, "normal": 11.6, "name": 21.0, "info": 11.0, "heading": 12.6, "section_before": 12, "line": 1.26, "bullet": 11.1, "small": 10.5, "entry": 11.3, "gap": 3.8},
    "maxfill": {"margin": 1.00, "normal": 12.0, "name": 22.0, "info": 11.4, "heading": 13.0, "section_before": 14, "line": 1.30, "bullet": 11.5, "small": 10.8, "entry": 11.8, "gap": 4.6},
    "maxfill1": {"margin": 1.03, "normal": 12.2, "name": 22.5, "info": 11.6, "heading": 13.2, "section_before": 15, "line": 1.32, "bullet": 11.7, "small": 11.0, "entry": 12.0, "gap": 5.0},
    "maxfill2": {"margin": 1.06, "normal": 12.4, "name": 23.0, "info": 11.8, "heading": 13.4, "section_before": 16, "line": 1.34, "bullet": 11.9, "small": 11.2, "entry": 12.2, "gap": 5.4},
    "maxfill3": {"margin": 1.08, "normal": 12.6, "name": 23.5, "info": 12.0, "heading": 13.6, "section_before": 17, "line": 1.36, "bullet": 12.1, "small": 11.4, "entry": 12.4, "gap": 5.8},
}


def lv(compactness: str, key: str) -> float:
    return float(LAYOUTS[compactness][key])


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_FALLBACK)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_FALLBACK)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_spacing(p, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_bottom_border(paragraph, color="4F81BD", size="6"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def configure_doc(doc: Document, compactness: str):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    margin = lv(compactness, "margin")
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(min(margin, 0.44))
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(lv(compactness, "normal"))
    normal.font.color.rgb = TEXT


def add_text_line(doc: Document, text: str, size=9, bold=False, align=None, after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, after=after)
    if bold:
        run = p.add_run(text)
        set_run_font(run, size, bold, TEXT)
    else:
        add_rich_text(p, text, size)
    return p


def add_header(doc: Document, basics: dict[str, Any], compactness: str):
    name = basics.get("name") or "姓名"
    title = basics.get("title") or basics.get("intention") or "求职意向：产品经理"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=1 if compactness == "ultra" else 2)
    run = p.add_run(name)
    set_run_font(run, lv(compactness, "name"), True, ACCENT)

    parts = [title]
    for key in ["phone", "email", "city"]:
        value = basics.get(key)
        if value:
            parts.append(str(value))
    for link in basics.get("links") or []:
        if link:
            parts.append(str(link))
    add_text_line(doc, "｜".join(parts), size=lv(compactness, "info"), align=WD_ALIGN_PARAGRAPH.CENTER, after=2)


def add_section_heading(doc: Document, title: str, compactness: str):
    p = doc.add_paragraph()
    set_para_spacing(p, before=lv(compactness, "section_before"), after=1)
    run = p.add_run(title)
    set_run_font(run, lv(compactness, "heading"), True, ACCENT)
    add_bottom_border(p)
    return p


def add_education(doc: Document, items: list[dict[str, Any]], compactness: str):
    if not items:
        return
    add_section_heading(doc, "教育背景", compactness)
    for edu in items:
        main_parts = [edu.get("school"), edu.get("major"), edu.get("degree"), edu.get("time")]
        line = "｜".join([str(x) for x in main_parts if x])
        if line:
            add_text_line(doc, line, size=lv(compactness, "entry"), bold=True, after=0)
        details = edu.get("details") or []
        if details:
            add_text_line(doc, "；".join([str(x) for x in details if x]), size=lv(compactness, "small"), after=0)


IMPORTANT_PATTERN = re.compile(
    r"(\d+(?:\.\d+)?\s*(?:\+|%|％|w|W|万|千|百|小时|分钟|天|周|月|年|人|次|个|条|篇)|"
    r"\d+(?:\.\d+)?/\d+(?:\.\d+)?|"
    r"(?:TOP|Top|top)\s*\d+|"
    r"(?:GMV|ROI|CTR|CVR|DAU|MAU|SQL|A/B|PRD|SOP|RAG|LLM|AI))(?:[。；;，,])?"
)


def add_rich_text(paragraph, text: str, size: float, *, base_bold: bool = False, emphasize_metrics: bool = True):
    pos = 0
    text = str(text)
    for match in IMPORTANT_PATTERN.finditer(text) if emphasize_metrics else []:
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size, base_bold, TEXT)
        run = paragraph.add_run(match.group(0))
        set_run_font(run, size, True, TEXT)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size, base_bold, TEXT)


def clean_bullet_text(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text)).strip().rstrip("。.")


def add_bullet(doc: Document, text: str, compactness: str):
    p = doc.add_paragraph(style=None)
    set_para_spacing(p, after=lv(compactness, "gap"), line=lv(compactness, "line"))
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.10)
    run = p.add_run("• ")
    set_run_font(run, lv(compactness, "bullet"), False, ACCENT)
    add_rich_text(p, clean_bullet_text(text), lv(compactness, "bullet"))


def add_entry(doc: Document, entry: dict[str, Any], compactness: str):
    heading = entry.get("heading")
    if heading:
        add_text_line(doc, str(heading), size=lv(compactness, "entry"), bold=True, after=0)
    summary = entry.get("summary")
    if summary:
        add_text_line(doc, str(summary), size=lv(compactness, "small"), after=0)
    for project in entry.get("projects") or []:
        name = project.get("name")
        if name:
            add_text_line(doc, str(name), size=lv(compactness, "bullet"), bold=True, after=0)
        ps = project.get("summary")
        if ps:
            add_text_line(doc, str(ps), size=lv(compactness, "small"), after=0)
        for bullet in project.get("bullets") or []:
            add_bullet(doc, bullet, compactness)
    for bullet in entry.get("bullets") or []:
        add_bullet(doc, bullet, compactness)


def add_sections(doc: Document, sections: list[dict[str, Any]], compactness: str):
    for sec in sections:
        entries = sec.get("entries") or []
        if not entries:
            continue
        add_section_heading(doc, str(sec.get("title") or "经历"), compactness)
        for entry in entries:
            add_entry(doc, entry, compactness)


def add_skills(doc: Document, skills: list[Any], compactness: str):
    if not skills:
        return
    add_section_heading(doc, "技能与其他", compactness)
    size = lv(compactness, "bullet")
    lines: list[str] = []
    for item in skills:
        if not item:
            continue
        if isinstance(item, dict):
            label = str(item.get("label") or item.get("category") or "").strip()
            values = item.get("items") or item.get("values") or item.get("text") or ""
            if isinstance(values, list):
                values = "、".join([str(v) for v in values if v])
            line = f"{label}：{values}" if label else str(values)
        else:
            line = str(item).strip()
        if line:
            # Split semicolon-packed skill rows into separate readable rows.
            parts = [p.strip() for p in re.split(r"[;；]", line) if p.strip()]
            lines.extend(parts if len(parts) > 1 else [line])
    for line in lines:
        p = doc.add_paragraph()
        set_para_spacing(p, after=lv(compactness, "gap"), line=lv(compactness, "line"))
        add_rich_text(p, line, size)


def scrub_metadata(doc: Document):
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.subject = ""
    props.keywords = ""
    props.title = ""


def estimate_content_lines(data: dict[str, Any]) -> int:
    lines = 2 + max(1, len(data.get("education") or []))
    for sec in data.get("sections") or []:
        entries = sec.get("entries") or []
        if entries:
            lines += 1
        for entry in entries:
            lines += 1
            if entry.get("summary"):
                lines += max(1, len(str(entry["summary"])) // 54 + 1)
            for project in entry.get("projects") or []:
                if project.get("name"):
                    lines += 1
                if project.get("summary"):
                    lines += max(1, len(str(project["summary"])) // 54 + 1)
                for bullet in project.get("bullets") or []:
                    lines += max(1, len(str(bullet)) // 48 + 1)
            for bullet in entry.get("bullets") or []:
                lines += max(1, len(str(bullet)) // 48 + 1)
    skill_count = 0
    for item in data.get("skills") or []:
        text = str(item.get("text") or item.get("items") or item) if isinstance(item, dict) else str(item)
        skill_count += max(1, len([p for p in re.split(r"[;；]", text) if p.strip()]))
    if skill_count:
        lines += 1 + skill_count
    return lines


def choose_compactness(data: dict[str, Any]) -> str:
    lines = estimate_content_lines(data)
    if lines <= 22:
        return "maxfill2"
    if lines <= 25:
        return "maxfill"
    if lines <= 28:
        return "fill4"
    if lines <= 31:
        return "fill3"
    if lines <= 34:
        return "fill2"
    if lines <= 37:
        return "fill1"
    if lines <= 42:
        return "expanded"
    if lines <= 50:
        return "roomy"
    if lines <= 62:
        return "normal"
    if lines <= 76:
        return "tight"
    return "ultra"


def build(data: dict[str, Any], output: Path, compactness: str):
    if compactness == "auto":
        compactness = choose_compactness(data)
    doc = Document()
    configure_doc(doc, compactness)
    add_header(doc, data.get("basics") or {}, compactness)
    add_education(doc, data.get("education") or [], compactness)
    add_sections(doc, data.get("sections") or [], compactness)
    add_skills(doc, data.get("skills") or [], compactness)
    scrub_metadata(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return compactness


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--compactness", choices=["auto", *LAYOUTS.keys()], default="auto")
    args = parser.parse_args()

    data = json.loads(args.input_json.read_text(encoding="utf-8"))
    compactness = build(data, args.output_docx, args.compactness)
    print(f"Wrote {args.output_docx} compactness={compactness}")


if __name__ == "__main__":
    main()

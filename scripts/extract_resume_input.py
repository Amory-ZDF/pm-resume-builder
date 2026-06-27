#!/usr/bin/env python3
"""Extract plain text from resume inputs: DOCX, PDF, TXT, or Markdown.

The script prints extracted text to stdout by default. Use --out to write a file.
It does not anonymize content; agents should treat output as private user data.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED = {".docx", ".pdf", ".txt", ".md", ".markdown"}


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    blank = False
    for raw in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if not line:
            if not blank:
                lines.append("")
            blank = True
            continue
        lines.append(line)
        blank = False
    return "\n".join(lines).strip() + "\n"


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", "；") for cell in row.cells]
            row_text = "｜".join([c for c in cells if c])
            if row_text:
                parts.append(row_text)
    return normalize_text("\n".join(parts))


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"[Page {i}]\n{text}")
    return normalize_text("\n\n".join(parts))


def extract_plain(path: Path) -> str:
    return normalize_text(path.read_text(encoding="utf-8", errors="replace"))


def extract(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return extract_plain(path)
    raise SystemExit(f"Unsupported file type: {suffix}. Supported: {', '.join(sorted(SUPPORTED))}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="Resume input file: .docx, .pdf, .txt, .md")
    parser.add_argument("--out", type=Path, help="Optional output text path")
    args = parser.parse_args()

    if not args.input.exists():
        raise SystemExit(f"Input file not found: {args.input}")
    text = extract(args.input)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(text, encoding="utf-8")
        print(f"Wrote {args.out}")
    else:
        print(text, end="")


if __name__ == "__main__":
    main()
